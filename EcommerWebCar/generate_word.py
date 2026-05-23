from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()
doc.add_heading('AI Chatbot Architecture & Implementation', 0)

doc.add_paragraph(
    "This document outlines the architecture and workflow of the AI Chatbot built for the Car Shop Microservice project. "
    "It is designed to help you prepare for your technical interview by explaining the exact technologies and processes used in the system."
)

doc.add_heading('1. Overview of the AI Chatbot', level=1)
doc.add_paragraph(
    "Instead of a simple chatbot that only reads from a static database, we built an 'AI Agent' using LangChain. "
    "This means the chatbot can 'think' and decide which tool to use based on the customer's question. "
    "The core brain of the chatbot is OpenAI's gpt-4o-mini model."
)

doc.add_heading('2. The Technologies Used', level=1)
p = doc.add_paragraph()
p.add_run('1. Large Language Model (LLM): ').bold = True
p.add_run('OpenAI (gpt-4o-mini) acts as the reasoning engine.\n')

p.add_run('2. Orchestration Framework: ').bold = True
p.add_run('LangChain is used to build the Tool-Calling Agent and connect all the pieces.\n')

p.add_run('3. Vector Database: ').bold = True
p.add_run('ChromaDB is used to store text embeddings for the RAG system.\n')

p.add_run('4. Embedding Model: ').bold = True
p.add_run('OpenAI (text-embedding-3-small) converts text into numerical vectors.\n')

p.add_run('5. Backend: ').bold = True
p.add_run('Python, running as a microservice.')

doc.add_heading('3. How the Tools Work', level=1)
doc.add_paragraph(
    "The AI Agent is equipped with 4 specialized tools. When a user asks a question, the Agent analyzes the intent and selects the best tool:"
)

p2 = doc.add_paragraph(style='List Bullet')
p2.add_run('RAG Tool (Car_FAQ_Knowledge): ').bold = True
p2.add_run(
    "This tool is used for general questions like warranties, car comparisons, or FAQs. "
    "It uses Retrieval-Augmented Generation (RAG). First, internal Markdown documents are split into small chunks. "
    "These chunks are embedded and saved into ChromaDB. When a customer asks a question, we search ChromaDB for the top 3 most relevant chunks and feed them to the LLM to generate a precise answer."
)

p3 = doc.add_paragraph(style='List Bullet')
p3.add_run('MySQL Query Tool: ').bold = True
p3.add_run(
    "This is used when the customer wants real-time data, like checking inventory or searching for cars by price. "
    "Instead of guessing, the Agent triggers this tool to securely query the showroom's actual MySQL database and return structured car cards to the frontend."
)

p4 = doc.add_paragraph(style='List Bullet')
p4.add_run('Rolling Price Calculator: ').bold = True
p4.add_run(
    "Used when the user asks 'how much does this car cost to drive home?'. "
    "It calculates taxes, registration fees, and other costs based on the customer's location (e.g., Ho Chi Minh City)."
)

p5 = doc.add_paragraph(style='List Bullet')
p5.add_run('Loan & Finance Guidance: ').bold = True
p5.add_run(
    "A fallback tool used when customers ask about bank loans or interest rates. It guides them to contact the showroom directly."
)

doc.add_heading('4. The Complete Workflow (Step-by-Step)', level=1)
doc.add_paragraph(
    "If an interviewer asks how a request is processed from start to finish, explain this flow:"
)
doc.add_paragraph("The customer sends a message (e.g., 'Do you have the Mazda CX-5 and how much is it?').", style='List Number')
doc.add_paragraph("The LangChain Agent receives the message and analyzes the intent.", style='List Number')
doc.add_paragraph("The Agent decides it needs real-time data, so it calls the 'MySQL Query Tool' with parameters like car_name='Mazda CX-5'.", style='List Number')
doc.add_paragraph("The tool executes the query and returns the live database results.", style='List Number')
doc.add_paragraph("The Agent combines this data into a friendly, natural language response and sends it back to the customer.", style='List Number')

doc.add_heading('5. Quality Assurance (Evaluation)', level=1)
doc.add_paragraph(
    "To ensure the chatbot provides accurate answers, we implemented an evaluation script using the RAGAS framework. "
    "This automatically tests the chatbot on metrics like 'Faithfulness' (not hallucinating information) and 'Answer Relevance' (staying on topic)."
)

doc.save('Chatbot_Interview_Guide.docx')
